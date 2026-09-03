"""md → docx 转换器（复赛报告排版版）。

处理：#/##/###/#### 标题、**粗体**、`代码`、| 表格 |、- /数字 列表、> 引用、--- 分隔线（跳过）。
排版：A4 适中边距、微软雅黑正文 11pt/1.35 倍行距、藏青标题层级、表格藏青表头+斑马纹。
用法: python -m src.md2docx <in.md> <out.docx>
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x4E, 0x79)
NAVY2 = RGBColor(0x2E, 0x5A, 0x88)
INK = RGBColor(0x22, 0x2A, 0x33)
GRAY = RGBColor(0x60, 0x60, 0x60)
CODE_RED = RGBColor(0xB0, 0x30, 0x60)


def add_runs(par, text):
    """处理 **粗体** 与 `代码` 行内标记。"""
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.color.rgb = CODE_RED
        else:
            par.add_run(tok)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_par_shading(par, hex_color):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_par_bottom_border(par, hex_color="1F4E79", sz=8):
    pPr = par._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(2.3)
        sec.bottom_margin = Cm(2.1)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(2.4)

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.font.color.rgb = INK
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    pf = st.paragraph_format
    pf.line_spacing = 1.35
    pf.space_after = Pt(5)

    specs = {1: (19, NAVY, 18, 8), 2: (15, NAVY, 14, 6),
             3: (12.5, NAVY2, 10, 4), 4: (11, NAVY2, 8, 3)}
    for i, (size, color, before, after) in specs.items():
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Calibri"
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = color
        hs.paragraph_format.space_before = Pt(before)
        hs.paragraph_format.space_after = Pt(after)
        hs.paragraph_format.line_spacing = 1.15


def flush_table(doc, rows):
    if not rows:
        return
    cells = [[c.strip() for c in re.split(r"(?<!\\)\|", r)[1:-1]] for r in rows]
    cells = [c for c in cells if not all(re.fullmatch(r":?-{2,}:?", x or "--")
                                         for x in c)]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(cells):
        for j in range(ncol):
            cell = t.cell(i, j)
            txt = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, txt.replace("\\|", "|"))
            for r in p.runs:
                r.font.size = Pt(9.5)
                if i == 0:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                shade_cell(cell, "1F4E79")
            elif i % 2 == 0:
                shade_cell(cell, "F2F5F8")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def convert(src, dst):
    import datetime
    doc = Document()
    setup(doc)
    # 元数据中性化（去掉 python-docx 默认工具痕迹；作者/描述留空，
    # 后续在 WPS 中编辑保存时会自然写入本机用户信息）
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.title = "虚拟酵母细胞_蛋白质组扰动响应的泛化预测（复赛报告）"
    cp.subject = ""
    cp.comments = ""
    cp.category = ""
    cp.created = datetime.datetime(2026, 9, 2, 12, 0, 0)
    cp.modified = datetime.datetime(2026, 9, 2, 12, 0, 0)

    lines = open(src, encoding="utf-8").read().split("\n")
    tbl = []
    for line in lines:
        if line.strip().startswith("|"):
            tbl.append(line)
            continue
        if tbl:
            flush_table(doc, tbl)
            tbl = []
        s = line.rstrip()
        if not s.strip() or re.fullmatch(r"-{3,}", s.strip()):
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", s)
        if m:
            level = len(m.group(1))
            doc.add_heading("", level=level)
            par = doc.paragraphs[-1]
            add_runs(par, m.group(2))
            if level == 2:
                set_par_bottom_border(par)
            continue
        if s.startswith(">"):
            par = doc.add_paragraph()
            add_runs(par, s.lstrip("> ").strip())
            par.paragraph_format.left_indent = Pt(14)
            set_par_shading(par, "F2F5F8")
            for r in par.runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = GRAY
            continue
        m = re.match(r"^(\s*)-\s+(.*)", s)
        if m:
            par = doc.add_paragraph(style="List Bullet")
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.line_spacing = 1.3
            add_runs(par, m.group(2))
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)", s)
        if m:
            par = doc.add_paragraph(style="List Number")
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.line_spacing = 1.3
            add_runs(par, m.group(2))
            continue
        par = doc.add_paragraph()
        add_runs(par, s)
    if tbl:
        flush_table(doc, tbl)
    doc.save(dst)
    print(f"[saved] {dst}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
